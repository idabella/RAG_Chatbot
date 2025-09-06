import mimetypes
import hashlib
import logging
from pathlib import Path
from typing import List, Dict, Any, Union, Optional
from datetime import datetime
import chardet

import PyPDF2
import docx
import csv 
import json
from fastapi import UploadFile

logger = logging.getLogger(__name__)


class FileValidationError(Exception):
    pass


class FileExtractionError(Exception):
    pass


class FileValidator:
    SUPPORTED_FORMATS = {
        'pdf': ['application/pdf'],
        'docx': ['application/vnd.openxmlformats-officedocument.wordprocessingml.document'],
        'txt': ['text/plain'],
        'md': ['text/markdown', 'text/x-markdown']
    }
    
    MAX_FILE_SIZE = 50 * 1024 * 1024
    
    DANGEROUS_EXTENSIONS = {
        '.exe', '.bat', '.cmd', '.com', '.pif', '.scr', '.vbs', '.js', '.jar',
        '.sh', '.php', '.asp', '.aspx', '.jsp', '.py', '.pl', '.rb'
    }

    @staticmethod
    def validate_file_format(file: UploadFile) -> bool:
        try:
            file_extension = Path(file.filename).suffix.lower()
            if file_extension.lstrip('.') not in FileValidator.SUPPORTED_FORMATS:
                raise FileValidationError(f"Format de fichier non supporte: {file_extension}")
            
            if file_extension in FileValidator.DANGEROUS_EXTENSIONS:
                raise FileValidationError(f"Extension dangereuse detectee: {file_extension}")
            
            return True
            
        except Exception as e:
            logger.error(f"Erreur lors de la validation du fichier {file.filename}: {e}")
            raise FileValidationError(f"Erreur de validation: {str(e)}")

    @staticmethod
    def validate_file_size(file: UploadFile) -> bool:
        file.file.seek(0, 2)
        file_size = file.file.tell()
        file.file.seek(0)
        
        if file_size > FileValidator.MAX_FILE_SIZE:
            raise FileValidationError(
                f"Fichier trop volumineux: {file_size} bytes "
                f"(max: {FileValidator.MAX_FILE_SIZE} bytes)"
            )
        
        return True

    @staticmethod
    def detect_file_type(file_path: Union[str, Path]) -> str:
        try:
            mime_type, _ = mimetypes.guess_type(str(file_path))
            return mime_type or 'application/octet-stream'
        except Exception as e:
            logger.warning(f"Impossible de detecter le type de fichier {file_path}: {e}")
            return 'application/octet-stream'

    @staticmethod
    def generate_file_hash(file_path: Union[str, Path]) -> str:
        hash_sha256 = hashlib.sha256()
        
        with open(file_path, 'rb') as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_sha256.update(chunk)
        
        return hash_sha256.hexdigest()

    @staticmethod
    def is_text_file(file_path: Union[str, Path]) -> bool:
        try:
            with open(file_path, 'rb') as f:
                sample = f.read(1024)
                result = chardet.detect(sample)
                return result['encoding'] is not None
        except Exception:
            return False


class FileExtractor:
    
    @staticmethod
    def extract_pdf(file_path: Union[str, Path]) -> str:
        text = ""
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            logger.error(f"Erreur extraction PDF: {e}")
            raise FileExtractionError(f"Erreur extraction PDF: {str(e)}")
        return text.strip()
    
    @staticmethod
    def extract_docx(file_path: Union[str, Path]) -> str:
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text += " | ".join(row_text) + "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"Erreur extraction DOCX: {e}")
            raise FileExtractionError(f"Erreur extraction DOCX: {str(e)}")
    
    @staticmethod
    def extract_txt(file_path: Union[str, Path], encoding: Optional[str] = None) -> str:
        try:
            if encoding:
                with open(file_path, 'r', encoding=encoding) as file:
                    return file.read()
            
            with open(file_path, 'rb') as file:
                raw_data = file.read()
                result = chardet.detect(raw_data)
                detected_encoding = result['encoding'] or "utf-8"
            
            with open(file_path, 'r', encoding=detected_encoding) as file:
                return file.read()
                
        except Exception as e:
            logger.error(f"Erreur extraction TXT: {e}")
            raise FileExtractionError(f"Erreur extraction TXT: {str(e)}")
    
    @staticmethod
    def extract_md(file_path: Union[str, Path]) -> str:
        try:
            content = FileExtractor.extract_txt(file_path)
            
            import re
            content = re.sub(r'^#+\s*', '', content, flags=re.MULTILINE)
            content = re.sub(r'\*{1,2}([^*]+)\*{1,2}', r'\1', content)
            content = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', content)
            content = re.sub(r'```[^`]*```', '', content, flags=re.DOTALL)
            content = re.sub(r'`([^`]+)`', r'\1', content)
            
            return content.strip()
        except Exception as e:
            logger.error(f"Erreur extraction MD: {e}")
            raise FileExtractionError(f"Erreur extraction MD: {str(e)}")
    
    @staticmethod
    def extract_content(file_path: Union[str, Path]) -> str:
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        extractors = {
            '.pdf': FileExtractor.extract_pdf,
            '.docx': FileExtractor.extract_docx,
            '.txt': FileExtractor.extract_txt,
            '.md': FileExtractor.extract_md,
        }
        
        if extension not in extractors:
            raise FileExtractionError(f"Type de fichier non supporte: {extension}")
        
        if not file_path.exists():
            raise FileNotFoundError(f"Fichier non trouve: {file_path}")
        
        return extractors[extension](file_path)


class FileUtils:
    
    @staticmethod
    async def save_upload_file(file: UploadFile, destination: Union[str, Path]) -> Path:
        try:
            destination = Path(destination)
            destination.parent.mkdir(parents=True, exist_ok=True)
            
            content = await file.read()
            with open(destination, 'wb') as f:
                f.write(content)
            
            await file.seek(0)
            
            logger.info(f"Fichier sauvegarde: {destination}")
            return destination
            
        except Exception as e:
            logger.error(f"Erreur lors de la sauvegarde de {file.filename}: {e}")
            raise FileValidationError(f"Erreur de sauvegarde: {str(e)}")
    
    @staticmethod
    def get_file_info(file_path: Union[str, Path]) -> Dict[str, Any]:
        try:
            file_path = Path(file_path)
            stat = file_path.stat()
            
            return {
                'name': file_path.name,
                'size': stat.st_size,
                'size_human': FileUtils.format_file_size(stat.st_size),
                'extension': file_path.suffix.lower(),
                'mime_type': FileValidator.detect_file_type(file_path),
                'created_at': datetime.fromtimestamp(stat.st_ctime),
                'modified_at': datetime.fromtimestamp(stat.st_mtime),
                'hash': FileValidator.generate_file_hash(file_path),
                'is_text': FileValidator.is_text_file(file_path),
                'is_supported': file_path.suffix.lower().lstrip('.') in FileValidator.SUPPORTED_FORMATS
            }
            
        except Exception as e:
            logger.error(f"Erreur lors de la recuperation des infos de {file_path}: {e}")
            return {}

    @staticmethod
    def format_file_size(size_bytes: int) -> str:
        for unit in ['B', 'KB', 'MB', 'GB', 'TB']:
            if size_bytes < 1024.0:
                return f"{size_bytes:.1f} {unit}"
            size_bytes /= 1024.0
        return f"{size_bytes:.1f} PB"
    
    @staticmethod
    def clean_filename(filename: str) -> str:
        import re
        filename = re.sub(r'[^\w\s\-_\.]', '', filename)
        filename = re.sub(r'[-\s]+', '-', filename)
        return filename.strip('-')
    
    @staticmethod
    def generate_unique_filename(original_filename: str, directory: Union[str, Path]) -> str:
        directory = Path(directory)
        base_name = Path(original_filename).stem
        extension = Path(original_filename).suffix
        
        clean_base = FileUtils.clean_filename(base_name)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        unique_filename = f"{timestamp}_{clean_base}{extension}"
        
        counter = 1
        while (directory / unique_filename).exists():
            unique_filename = f"{timestamp}_{clean_base}_{counter}{extension}"
            counter += 1
        
        return unique_filename
    
    @staticmethod
    def delete_file_safe(file_path: Union[str, Path]) -> bool:
        try:
            file_path = Path(file_path)
            if file_path.exists():
                file_path.unlink()
                logger.info(f"Fichier supprime: {file_path}")
                return True
            return False
        except Exception as e:
            logger.error(f"Erreur lors de la suppression de {file_path}: {e}")
            return False
    
    @staticmethod
    def create_backup(file_path: Union[str, Path]) -> Optional[Path]:
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                return None
            
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            backup_name = f"{file_path.stem}_backup_{timestamp}{file_path.suffix}"
            backup_path = file_path.parent / backup_name
            
            import shutil
            shutil.copy2(file_path, backup_path)
            
            logger.info(f"Sauvegarde creee: {backup_path}")
            return backup_path
            
        except Exception as e:
            logger.error(f"Erreur lors de la creation de la sauvegarde: {e}")
            return None


def validate_uploaded_file(file: UploadFile) -> bool:
    FileValidator.validate_file_format(file)
    FileValidator.validate_file_size(file)
    return True


async def process_uploaded_file(file: UploadFile, 
                              save_directory: Union[str, Path]) -> Dict[str, Any]:
    try:
        validate_uploaded_file(file)
        
        save_directory = Path(save_directory)
        filename = FileUtils.generate_unique_filename(file.filename, save_directory)
        file_path = save_directory / filename
        
        await FileUtils.save_upload_file(file, file_path)
        content = FileExtractor.extract_content(file_path)
        
        file_info = FileUtils.get_file_info(file_path)
        file_info.update({
            'original_filename': file.filename,
            'saved_path': str(file_path),
            'content': content,
            'content_length': len(content),
            'processed_at': datetime.now()
        })
        
        logger.info(f"Fichier traite avec succès: {file.filename}")
        return file_info
        
    except Exception as e:
        logger.error(f"Erreur lors du traitement de {file.filename}: {e}")
        raise FileValidationError(f"Erreur de traitement: {str(e)}")


def get_file_content(file_path: Union[str, Path]) -> str:
    return FileExtractor.extract_content(file_path)

