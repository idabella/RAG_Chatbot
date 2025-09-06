import os
import hashlib
import mimetypes
from typing import List, Dict, Optional, Tuple, Any
from pathlib import Path
from datetime import datetime
import logging
import re
import asyncio
from concurrent.futures import ThreadPoolExecutor

import PyPDF2
import docx
import chardet
from pydantic import BaseModel

# Ajout d'imports pour une meilleure extraction PDF
try:
    import fitz  # PyMuPDF - plus robuste
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    
try:
    import pdfplumber  # Alternative pour les tableaux
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

from sqlalchemy.orm import Session
from sqlalchemy import and_, or_, desc

from models.document import Document, DocumentChunk, DocumentType, EmbeddingStatus
from core.config import settings
from utils.text_processing import TextProcessor
from utils.file_utils import FileValidator
from services.embedding_service import EmbeddingService

logger = logging.getLogger(__name__)


class DocumentMetadata(BaseModel):
    filename: str
    file_size: int
    file_type: str
    mime_type: str
    encoding: Optional[str] = None
    page_count: Optional[int] = None
    word_count: Optional[int] = None
    language: Optional[str] = None
    checksum: str
    created_at: datetime
    processed_at: Optional[datetime] = None


class DocumentChunkData(BaseModel):
    content: str
    chunk_index: int
    start_position: int
    end_position: int
    metadata: Dict[str, Any]


class DocumentProcessingResult(BaseModel):
    document_id: int
    success: bool
    chunks_count: int
    error_message: Optional[str] = None
    processing_time: float
    metadata: Optional[DocumentMetadata] = None


class DocumentService:    
    def __init__(self, db: Session, embedding_service: EmbeddingService = None):
        self.db = db
        self.embedding_service = embedding_service
        self.text_processor = TextProcessor()
        self.file_validator = FileValidator()
        self.executor = ThreadPoolExecutor(max_workers=getattr(settings, 'MAX_CONCURRENT_UPLOADS', 3))
        
        self.chunk_size = getattr(settings, 'RAG_CHUNK_SIZE', 1000)
        self.chunk_overlap = getattr(settings, 'RAG_CHUNK_OVERLAP', 200)
        
        self.supported_formats = {
            '.pdf': self._extract_pdf_content,
            '.docx': self._extract_docx_content,
            '.txt': self._extract_txt_content,
            '.md': self._extract_txt_content
        }

    async def process_document(
        self, 
        file_path: str, 
        user_id: int,
        category: Optional[str] = None,
        tags: Optional[List[str]] = None
    ) -> DocumentProcessingResult:
        """Traiter un document et l'indexer"""
        start_time = datetime.now()
        metadata = None
        
        try:
            # Validation du fichier
            if not await self._validate_file(file_path):
                raise ValueError(f"Fichier invalide: {file_path}")
            
            # Extraction des métadonnées
            metadata = await self._extract_metadata(file_path)
            
            # Vérification des doublons
            existing_doc = self._check_duplicate(metadata.checksum)
            if existing_doc:
                logger.warning(f"Document déjà existant: {metadata.filename}")
                return DocumentProcessingResult(
                    document_id=existing_doc.id,
                    success=False,
                    chunks_count=0,
                    error_message="Document déjà existant",
                    processing_time=0.0,
                    metadata=metadata
                )
            
            # Extraction du contenu
            content = await self._extract_content(file_path, metadata.file_type)
            
            # Prétraitement du texte
            cleaned_content = await self._preprocess_text(content)
            
            # Chunking amélioré avec extraction d'informations personnelles
            chunks = await self._chunk_text_with_personal_info(cleaned_content, metadata)
            
            # Création du document en base
            document = await self._create_document(
                file_path, metadata, user_id, category, tags, len(chunks)
            )
            
            # Sauvegarde des chunks
            await self._save_chunks(document.id, chunks)
            
            # ✅ CORRECTION: Générer les embeddings IMMÉDIATEMENT après la sauvegarde
            # sans nettoyer le service d'embeddings
            if self.embedding_service:
                try:
                    await self._generate_embeddings_sync(document.id, chunks)
                    logger.info(f"Embeddings générés avec succès pour le document {document.id}")
                except Exception as embed_error:
                    logger.error(f"Erreur génération embeddings: {str(embed_error)}")
                    # Marquer le document comme échoué
                    document.embeddings_status = EmbeddingStatus.FAILED
                    self.db.commit()
            
            # Commit final
            self.db.commit()
            
            processing_time = (datetime.now() - start_time).total_seconds()
            
            return DocumentProcessingResult(
                document_id=document.id,
                success=True,
                chunks_count=len(chunks),
                processing_time=processing_time,
                metadata=metadata
            )
            
        except Exception as e:
            logger.error(f"Erreur lors du traitement du document {file_path}: {str(e)}")
            self.db.rollback()
            processing_time = (datetime.now() - start_time).total_seconds()
            
            return DocumentProcessingResult(
                document_id=0,
                success=False,
                chunks_count=0,
                error_message=str(e),
                processing_time=processing_time,
                metadata=metadata
            )

    async def _generate_embeddings_sync(self, document_id: int, chunks: List[DocumentChunkData]):
        """Générer les embeddings de façon synchrone - NOUVELLE MÉTHODE CORRIGÉE"""
        try:
            # Récupérer le document
            document = self.db.query(Document).filter(Document.id == document_id).first()
            if not document:
                logger.error(f"Document {document_id} non trouvé pour génération embeddings")
                return
            
            # Marquer comme en cours de traitement
            document.embeddings_status = EmbeddingStatus.PROCESSING
            self.db.commit()
            
            logger.info(f"Début de l'indexation pour le document {document_id} avec {len(chunks)} chunks")
            
            # ✅ CORRECTION: Vérifier que le service d'embeddings est bien initialisé
            if not self.embedding_service:
                logger.error("Service d'embeddings non disponible")
                document.embeddings_status = EmbeddingStatus.FAILED
                self.db.commit()
                return
            
            # ✅ CORRECTION: Vérifier l'état interne du service d'embeddings
            try:
                # Réinitialiser le service si nécessaire
                if not hasattr(self.embedding_service, 'model') or not self.embedding_service.model:
                    logger.info("Réinitialisation du service d'embeddings...")
                    await self.embedding_service._initialize_model()
                    await self.embedding_service._initialize_chroma()
                
                if not hasattr(self.embedding_service, 'collection') or not self.embedding_service.collection:
                    logger.info("Réinitialisation de la collection ChromaDB...")
                    await self.embedding_service._initialize_chroma()
                
            except Exception as init_error:
                logger.error(f"Erreur réinitialisation service embeddings: {str(init_error)}")
                document.embeddings_status = EmbeddingStatus.FAILED
                self.db.commit()
                return
            
            # Générer les embeddings pour chaque chunk
            successful_embeddings = 0
            for chunk_data in chunks:
                try:
                    await self.embedding_service.index_document(
                        document_id=str(document_id),
                        content=chunk_data.content,
                        metadata={
                            "document_id": document_id,
                            "source_file": document.filename,
                            "chunk_index": chunk_data.chunk_index,
                            **chunk_data.metadata  # Inclure les metadata enrichies
                        }
                    )
                    successful_embeddings += 1
                    logger.debug(f"Embedding généré pour chunk {chunk_data.chunk_index}")
                    
                except Exception as embed_error:
                    logger.error(f"Erreur embedding chunk {chunk_data.chunk_index}: {str(embed_error)}")
                    continue
            
            # Vérifier le succès
            if successful_embeddings == 0:
                logger.error(f"Aucun embedding généré pour le document {document_id}")
                document.embeddings_status = EmbeddingStatus.FAILED
            elif successful_embeddings < len(chunks):
                logger.warning(f"Seulement {successful_embeddings}/{len(chunks)} embeddings générés")
                document.embeddings_status = EmbeddingStatus.COMPLETED  # Marquer comme complété même partiellement
            else:
                logger.info(f"Tous les embeddings générés avec succès ({successful_embeddings}/{len(chunks)})")
                document.embeddings_status = EmbeddingStatus.COMPLETED
            
            # Sauvegarder le statut
            self.db.commit()
            
        except Exception as e:
            logger.error(f"Erreur génération embeddings pour document {document_id}: {str(e)}")
            try:
                document = self.db.query(Document).filter(Document.id == document_id).first()
                if document:
                    document.embeddings_status = EmbeddingStatus.FAILED
                    self.db.commit()
            except:
                pass

    async def _validate_file(self, file_path: str) -> bool:
        """Valider un fichier"""
        try:
            path = Path(file_path)
            
            if not path.exists():
                logger.error(f"Fichier inexistant: {file_path}")
                return False
            
            max_file_size = getattr(settings, 'MAX_FILE_SIZE', 10 * 1024 * 1024)  # 10MB par défaut
            if path.stat().st_size > max_file_size:
                logger.error(f"Fichier trop volumineux: {file_path}")
                return False
            
            if path.suffix.lower() not in self.supported_formats:
                logger.error(f"Format non supporté: {path.suffix}")
                return False
            
            return True
            
        except Exception as e:
            logger.error(f"Erreur validation fichier: {str(e)}")
            return False

    async def _extract_metadata(self, file_path: str) -> DocumentMetadata:
        """Extraire les métadonnées d'un fichier"""
        try:
            path = Path(file_path)
            stat = path.stat()
            
            checksum = await self._calculate_checksum(file_path)
            mime_type, _ = mimetypes.guess_type(file_path)
            
            encoding = None
            if path.suffix.lower() in ['.txt', '.md']:
                encoding = await self._detect_encoding(file_path)
            
            return DocumentMetadata(
                filename=path.name,
                file_size=stat.st_size,
                file_type=path.suffix.lower(),
                mime_type=mime_type or "application/octet-stream",
                encoding=encoding,
                checksum=checksum,
                created_at=datetime.fromtimestamp(stat.st_ctime)
            )
        except Exception as e:
            logger.error(f"Erreur extraction métadonnées: {str(e)}")
            raise

    async def _calculate_checksum(self, file_path: str) -> str:
        """Calculer le checksum MD5 d'un fichier"""
        def _hash_file():
            hash_md5 = hashlib.md5()
            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_md5.update(chunk)
            return hash_md5.hexdigest()
        
        try:
            loop = asyncio.get_event_loop()
            return await loop.run_in_executor(self.executor, _hash_file)
        except Exception as e:
            logger.error(f"Erreur calcul checksum: {str(e)}")
            raise

    async def _detect_encoding(self, file_path: str) -> str:
        """Détecter l'encodage d'un fichier texte"""
        def _detect():
            with open(file_path, 'rb') as f:
                raw_data = f.read(10000)
                result = chardet.detect(raw_data)
                return result['encoding'] or "utf-8"
        
        try:
            loop = asyncio.get_event_loop()
            return await loop.run_in_executor(self.executor, _detect)
        except Exception as e:
            logger.error(f"Erreur détection encodage: {str(e)}")
            return "utf-8"

    async def _extract_content(self, file_path: str, file_type: str) -> str:
        """Extraire le contenu d'un fichier"""
        extractor = self.supported_formats.get(file_type)
        if not extractor:
            raise ValueError(f"Format non supporté: {file_type}")
        
        try:
            loop = asyncio.get_event_loop()
            return await loop.run_in_executor(self.executor, extractor, file_path)
        except Exception as e:
            logger.error(f"Erreur extraction contenu: {str(e)}")
            raise

    def _extract_pdf_content(self, file_path: str) -> str:
        """Extraction PDF améliorée avec plusieurs méthodes"""
        content_parts = []
        
        # Méthode 1: PyMuPDF (fitz) - plus robuste pour les noms et caractères spéciaux
        if PYMUPDF_AVAILABLE:
            try:
                doc = fitz.open(file_path)
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    text = page.get_text()
                    if text.strip():
                        content_parts.append(f"[Page {page_num + 1}]\n{text}")
                doc.close()
                
                if content_parts:
                    logger.info(f"Extraction PyMuPDF réussie: {len(content_parts)} pages")
                    return "\n\n".join(content_parts)
            except Exception as e:
                logger.warning(f"Erreur PyMuPDF: {str(e)}")
        
        # Méthode 2: pdfplumber (fallback)
        if PDFPLUMBER_AVAILABLE:
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    content_parts = []
                    for page_num, page in enumerate(pdf.pages):
                        text = page.extract_text()
                        if text and text.strip():
                            content_parts.append(f"[Page {page_num + 1}]\n{text}")
                        
                        # Extraire aussi les tableaux
                        tables = page.extract_tables()
                        for table in tables:
                            table_text = "\n".join([" | ".join([str(cell or '') for cell in row]) for row in table])
                            if table_text.strip():
                                content_parts.append(f"[Tableau Page {page_num + 1}]\n{table_text}")
                
                if content_parts:
                    logger.info(f"Extraction pdfplumber réussie: {len(content_parts)} pages")
                    return "\n\n".join(content_parts)
            except Exception as e:
                logger.warning(f"Erreur pdfplumber: {str(e)}")
        
        # Méthode 3: PyPDF2 (fallback final)
        try:
            content = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if text.strip():
                            content.append(f"[Page {page_num + 1}]\n{text}")
                    except Exception as e:
                        logger.warning(f"Erreur extraction page {page_num + 1}: {str(e)}")
                        continue
            
            if content:
                logger.info(f"Extraction PyPDF2 réussie: {len(content)} pages")
                return "\n\n".join(content)
            
        except Exception as e:
            logger.error(f"Erreur PyPDF2: {str(e)}")
        
        raise ValueError(f"Impossible d'extraire le contenu du PDF: {file_path}")

    def _extract_docx_content(self, file_path: str) -> str:
        """Extraire le contenu d'un DOCX"""
        try:
            doc = docx.Document(file_path)
            content = []
            
            # Extraire les paragraphes
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    content.append(text)
            
            # Extraire les tableaux
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        content.append(" | ".join(row_text))
            
            return "\n\n".join(content)
            
        except Exception as e:
            logger.error(f"Erreur extraction DOCX: {str(e)}")
            raise

    def _extract_txt_content(self, file_path: str) -> str:
        """Extraire le contenu d'un fichier texte"""
        try:
            # Détecter l'encodage
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                encoding = chardet.detect(raw_data)['encoding'] or "utf-8"
            
            # Lire le contenu
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read()
                
        except Exception as e:
            logger.error(f"Erreur extraction TXT: {str(e)}")
            raise

    def _extract_personal_info(self, content: str) -> Dict[str, Any]:
        """Extraire les informations personnelles du CV - AMÉLIORATION CLÉE"""
        personal_info = {}
        
        # Patterns pour détecter les noms (améliorés pour ENSA Agadir)
        name_patterns = [
            # Patterns français
            r"^([A-ZÀ-Ÿ][a-zà-ÿ]+(?:\s+[A-ZÀ-Ÿ][a-zà-ÿ]+){1,3})(?:\s*$|\s+[^A-Za-z])",  # Nom en début de ligne
            r"Nom\s*:?\s*([A-ZÀ-Ÿ][a-zà-ÿ]+(?:\s+[A-ZÀ-Ÿ][a-zà-ÿ]+)+)",
            r"Name\s*:?\s*([A-ZÀ-Ÿ][a-zà-ÿ]+(?:\s+[A-ZÀ-Ÿ][a-zà-ÿ]+)+)",
            r"Prénom\s+et\s+nom\s*:?\s*([A-ZÀ-Ÿ][a-zà-ÿ]+(?:\s+[A-ZÀ-Ÿ][a-zà-ÿ]+)+)",
            r"Étudiant\s*:?\s*([A-ZÀ-Ÿ][a-zà-ÿ]+(?:\s+[A-ZÀ-Ÿ][a-zà-ÿ]+)+)",
            # Patterns arabes
            r"الاسم\s*:?\s*([^\n\r]+)",
            r"اسم\s*:?\s*([^\n\r]+)",
            # Pattern pour CV formatés
            r"^([A-Z][a-z]+\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)$",
        ]
        
        # Patterns pour l'école (spécifique ENSA Agadir)
        school_patterns = [
            r"ENSA\s+Agadir",
            r"École\s+Nationale\s+des\s+Sciences\s+Appliquées?\s+(?:d'|de\s+)?Agadir",
            r"National\s+School\s+of\s+Applied\s+Sciences?\s+Agadir",
            r"Agadir.*ENSA",
            r"ENSA.*Agadir",
        ]
        
        # Patterns pour la spécialisation
        specialty_patterns = [
            r"Data\s+Science",
            r"Science\s+des\s+Données",
            r"Sciences?\s+des\s+Données?",
            r"Ingénieur\s+Data\s+Science",
            r"Ingénierie\s+des\s+Données",
            r"Big\s+Data",
            r"Intelligence\s+Artificielle",
            r"Machine\s+Learning",
        ]
        
        lines = content.split('\n')
        
        # Chercher le nom dans les premières lignes (généralement en haut du CV)
        for i, line in enumerate(lines[:15]):  # Augmenté à 15 lignes
            line = line.strip()
            if not line or len(line) < 3:
                continue
            
            # Ignorer les lignes qui sont clairement des titres de section
            skip_patterns = [
                r"^(CV|CURRICULUM|VITAE|RESUME|CONTACT|FORMATION|EDUCATION|EXPERIENCE|COMPETENCE|SKILL)s?",
                r"^\d+",  # Lignes qui commencent par un chiffre
                r"^[@+\-=]",  # Lignes avec symboles
                r"^[0-9]{2}/[0-9]{2}/[0-9]{4}",  # Dates
                r"^Tel:",
                r"^Email:",
                r"^Page\s+\d+",
            ]
            
            skip_line = False
            for skip_pattern in skip_patterns:
                if re.search(skip_pattern, line, re.IGNORECASE):
                    skip_line = True
                    break
            
            if skip_line:
                continue
                
            for pattern in name_patterns:
                match = re.search(pattern, line, re.IGNORECASE)
                if match:
                    name = match.group(1).strip()
                    # Validation du nom
                    if self._is_valid_name(name):
                        personal_info['name'] = name
                        personal_info['name_line'] = i + 1
                        logger.info(f"Nom trouvé ligne {i+1}: {name}")
                        break
            
            if 'name' in personal_info:
                break
        
        # Chercher l'école dans tout le document
        full_content = content.lower()
        for pattern in school_patterns:
            match = re.search(pattern, content, re.IGNORECASE)
            if match:
                personal_info['school'] = match.group(0).strip()
                personal_info['school_confirmed'] = "ENSA Agadir"
                logger.info(f"École trouvée: {personal_info['school']}")
                break
        
        # Chercher la spécialisation
        for pattern in specialty_patterns:
            match = re.search(pattern, content, re.IGNORECASE)
            if match:
                personal_info['specialty'] = match.group(0).strip()
                if "data" in match.group(0).lower():
                    personal_info['specialty_confirmed'] = "Data Science"
                logger.info(f"Spécialité trouvée: {personal_info['specialty']}")
                break
        
        return personal_info

    def _is_valid_name(self, name: str) -> bool:
        """Valider qu'une chaîne est bien un nom de personne"""
        if not name or len(name) < 3:
            return False
        
        # Doit contenir au moins un espace (prénom + nom)
        if ' ' not in name:
            return False
        
        # Ne doit pas être trop long
        if len(name) > 50:
            return False
        
        # Ne doit pas contenir de chiffres ou symboles bizarres
        if re.search(r'[0-9@#$%^&*()_+=\[\]{}|\\:";\'<>?,./]', name):
            return False
        
        # Ne doit pas être une phrase complète
        words = name.split()
        if len(words) > 4:  # Maximum 4 mots pour un nom
            return False
        
        # Chaque mot doit commencer par une majuscule
        for word in words:
            if not word[0].isupper():
                return False
        
        # Mots interdits qui indiquent que ce n'est pas un nom
        forbidden_words = ['curriculum', 'vitae', 'resume', 'contact', 'formation', 
                          'education', 'experience', 'competence', 'skill', 'projet',
                          'stage', 'telephone', 'email', 'adresse', 'date']
        
        for word in words:
            if word.lower() in forbidden_words:
                return False
        
        return True

    async def _preprocess_text(self, content: str) -> str:
        """Prétraiter le texte"""
        try:
            if hasattr(self.text_processor, 'clean_text'):
                return self.text_processor.clean_text(content)
            else:
                # Nettoyage basique amélioré
                # Normaliser les espaces
                content = re.sub(r'\s+', ' ', content.strip())
                # Normaliser les retours à la ligne
                content = re.sub(r'\n\s*\n', '\n\n', content)
                # Supprimer les caractères de contrôle
                content = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', content)
                return content
        except Exception as e:
            logger.error(f"Erreur preprocessing: {str(e)}")
            return content

    async def _chunk_text_with_personal_info(self, content: str, metadata: DocumentMetadata) -> List[DocumentChunkData]:
        """Chunking amélioré qui préserve les informations personnelles - AMÉLIORATION CRUCIALE"""
        
        try:
            # Extraire les informations personnelles
            personal_info = self._extract_personal_info(content)
            
            chunks = []
            
            # Créer un chunk spécial avec les informations personnelles en premier
            if personal_info:
                info_lines = ["=== INFORMATIONS PERSONNELLES ==="]
                
                if 'name' in personal_info:
                    info_lines.append(f"Nom de l'étudiant: {personal_info['name']}")
                    info_lines.append(f"Étudiant ingénieur: {personal_info['name']}")
                
                if 'school_confirmed' in personal_info:
                    info_lines.append(f"École: {personal_info['school_confirmed']}")
                    info_lines.append(f"Institution: École Nationale des Sciences Appliquées d'Agadir")
                elif 'school' in personal_info:
                    info_lines.append(f"École: {personal_info['school']}")
                
                if 'specialty_confirmed' in personal_info:
                    info_lines.append(f"Spécialisation: {personal_info['specialty_confirmed']}")
                    info_lines.append(f"Domaine d'études: Sciences des Données")
                elif 'specialty' in personal_info:
                    info_lines.append(f"Spécialisation: {personal_info['specialty']}")
                
                # Ajouter une phrase de résumé pour les requêtes
                if 'name' in personal_info:
                    summary = f"L'étudiant ingénieur {personal_info['name']}"
                    if 'school_confirmed' in personal_info:
                        summary += f" étudie à {personal_info['school_confirmed']}"
                    if 'specialty_confirmed' in personal_info:
                        summary += f" en spécialisation {personal_info['specialty_confirmed']}"
                    summary += "."
                    info_lines.append(summary)
                
                info_text = "\n".join(info_lines)
                
                # Créer le chunk d'informations personnelles
                personal_chunk = DocumentChunkData(
                    content=info_text,
                    chunk_index=0,
                    start_position=0,
                    end_position=len(info_text),
                    metadata={
                        "filename": metadata.filename,
                        "file_type": metadata.file_type,
                        "chunk_type": "personal_info",
                        "chunk_length": len(info_text),
                        "word_count": len(info_text.split()),
                        "personal_info": personal_info,
                        "created_at": datetime.now().isoformat()
                    }
                )
                chunks.append(personal_chunk)
                
                logger.info(f"Chunk d'informations personnelles créé: {len(info_text)} caractères")
            
            # Chunking normal pour le reste du contenu
            regular_chunks = await self._chunk_text_enhanced(content, metadata, personal_info)
            
            # Ajouter les chunks réguliers en ajustant les index
            for chunk in regular_chunks:
                chunk.chunk_index = len(chunks)
                chunks.append(chunk)
            
            logger.info(f"Total de {len(chunks)} chunks créés (1 personnel + {len(regular_chunks)} réguliers)")
            
            return chunks
            
        except Exception as e:
            logger.error(f"Erreur chunking avec infos personnelles: {str(e)}")
            # Fallback vers le chunking simple
            return await self._chunk_text_simple(content, metadata)

    async def _chunk_text_enhanced(self, content: str, metadata: DocumentMetadata, personal_info: Dict[str, Any]) -> List[DocumentChunkData]:
        """Chunking amélioré qui enrichit chaque chunk avec les infos personnelles"""
        
        try:
            chunks = []
            paragraphs = content.split('\n\n')
            current_chunk = ""
            chunk_index = 0
            start_pos = 0
            
            # Préfixe à ajouter à chaque chunk pour améliorer la recherche
            context_prefix = ""
            if personal_info:
                prefix_parts = []
                if 'name' in personal_info:
                    prefix_parts.append(f"[Étudiant: {personal_info['name']}]")
                if 'school_confirmed' in personal_info:
                    prefix_parts.append(f"[École: {personal_info['school_confirmed']}]")
                if 'specialty_confirmed' in personal_info:
                    prefix_parts.append(f"[Spécialité: {personal_info['specialty_confirmed']}]")
                
                if prefix_parts:
                    context_prefix = " ".join(prefix_parts) + " "
            
            for paragraph in paragraphs:
                paragraph = paragraph.strip()
                if not paragraph:
                    continue
                
                # Paragraphe trop long
                if len(paragraph) > self.chunk_size:
                    # Sauvegarder le chunk actuel
                    if current_chunk:
                        enhanced_content = context_prefix + current_chunk
                        chunks.append(self._create_chunk_data(
                            enhanced_content, chunk_index, start_pos, 
                            start_pos + len(enhanced_content), metadata, personal_info
                        ))
                        chunk_index += 1
                        start_pos += len(enhanced_content)
                        current_chunk = ""
                    
                    # Diviser le paragraphe long
                    sub_chunks = self._split_long_paragraph(paragraph)
                    for sub_chunk in sub_chunks:
                        enhanced_content = context_prefix + sub_chunk
                        chunks.append(self._create_chunk_data(
                            enhanced_content, chunk_index, start_pos,
                            start_pos + len(enhanced_content), metadata, personal_info
                        ))
                        chunk_index += 1
                        start_pos += len(enhanced_content)
                
                # Le paragraphe peut être ajouté au chunk actuel
                elif len(current_chunk) + len(paragraph) + 2 <= self.chunk_size - len(context_prefix):
                    if current_chunk:
                        current_chunk += "\n\n" + paragraph
                    else:
                        current_chunk = paragraph
                
                # Le chunk actuel est plein, commencer un nouveau
                else:
                    if current_chunk:
                        enhanced_content = context_prefix + current_chunk
                        chunks.append(self._create_chunk_data(
                            enhanced_content, chunk_index, start_pos,
                            start_pos + len(enhanced_content), metadata, personal_info
                        ))
                        chunk_index += 1
                        start_pos += len(enhanced_content)
                    
                    current_chunk = paragraph
            
            # Ajouter le dernier chunk
            if current_chunk:
                enhanced_content = context_prefix + current_chunk
                chunks.append(self._create_chunk_data(
                    enhanced_content, chunk_index, start_pos,
                    start_pos + len(enhanced_content), metadata, personal_info
                ))
            
            return chunks
            
        except Exception as e:
            logger.error(f"Erreur chunking amélioré: {str(e)}")
            return await self._chunk_text_simple(content, metadata)

    def _create_chunk_data(
        self, 
        content: str, 
        index: int, 
        start: int, 
        end: int, 
        metadata: DocumentMetadata,
        personal_info: Dict[str, Any] = None
    ) -> DocumentChunkData:
        """Créer un objet DocumentChunkData enrichi"""
        
        chunk_metadata = {
            "filename": metadata.filename,
            "file_type": metadata.file_type,
            "chunk_length": len(content),
            "word_count": len(content.split()),
            "created_at": datetime.now().isoformat()
        }
        
        # Ajouter les informations personnelles au metadata du chunk
        if personal_info:
            chunk_metadata.update({
                "student_name": personal_info.get('name'),
                "school": personal_info.get('school_confirmed', personal_info.get('school')),
                "specialty": personal_info.get('specialty_confirmed', personal_info.get('specialty')),
                "has_personal_info": True
            })
        
        return DocumentChunkData(
            content=content,
            chunk_index=index,
            start_position=start,
            end_position=end,
            metadata=chunk_metadata
        )

    def _split_long_paragraph(self, paragraph: str) -> List[str]:
        """Diviser un paragraphe trop long en phrases"""
        chunks = []
        sentences = re.split(r'[.!?]+', paragraph)
        current_chunk = ""
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
                
            if len(current_chunk) + len(sentence) + 2 <= self.chunk_size:
                if current_chunk:
                    current_chunk += ". " + sentence
                else:
                    current_chunk = sentence
            else:
                if current_chunk:
                    chunks.append(current_chunk + ".")
                current_chunk = sentence
        
        if current_chunk:
            chunks.append(current_chunk + ".")
        
        return chunks

    async def _chunk_text_simple(self, content: str, metadata: DocumentMetadata) -> List[DocumentChunkData]:
        """Méthode de chunking simple par mots (fallback)"""
        chunks = []
        words = content.split()
        current_chunk = []
        current_length = 0
        chunk_index = 0
        
        for word in words:
            word_length = len(word) + 1
            
            if current_length + word_length > self.chunk_size and current_chunk:
                chunk_content = " ".join(current_chunk)
                chunks.append(self._create_chunk_data(
                    chunk_content, chunk_index, 0, len(chunk_content), metadata
                ))
                current_chunk = [word]
                current_length = word_length
                chunk_index += 1
            else:
                current_chunk.append(word)
                current_length += word_length
        
        if current_chunk:
            chunk_content = " ".join(current_chunk)
            chunks.append(self._create_chunk_data(
                chunk_content, chunk_index, 0, len(chunk_content), metadata
            ))
        
        return chunks

    def _check_duplicate(self, checksum: str) -> Optional[Document]:
        """Vérifier si un document existe déjà"""
        try:
            return self.db.query(Document).filter(Document.file_hash == checksum).first()
        except Exception as e:
            logger.error(f"Erreur vérification doublon: {str(e)}")
            return None

    async def _create_document(
        self, 
        file_path: str,
        metadata: DocumentMetadata, 
        user_id: int,
        category: Optional[str],
        tags: Optional[List[str]],
        chunk_count: int
    ) -> Document:
        """Créer un document en base de données"""
        try:
            document_type = self._get_document_type(metadata.file_type)
            
            document = Document(
                filename=metadata.filename,
                original_filename=metadata.filename,
                file_path=file_path,
                file_size=metadata.file_size,
                file_hash=metadata.checksum,
                document_type=document_type,
                mime_type=metadata.mime_type,
                embeddings_status=EmbeddingStatus.PENDING,
                category=category,
                tags=tags,
                word_count=metadata.word_count,
                chunk_count=chunk_count,
                uploaded_by=user_id
            )
            
            self.db.add(document)
            self.db.flush()  # Pour obtenir l'ID
            
            logger.info(f"Document créé avec ID: {document.id}")
            return document
            
        except Exception as e:
            logger.error(f"Erreur création document: {str(e)}")
            raise

    def _get_document_type(self, file_extension: str) -> DocumentType:
        """Mapper l'extension au type de document"""
        mapping = {
            '.pdf': DocumentType.PDF,
            '.docx': DocumentType.DOCX,
            '.txt': DocumentType.TXT,
            '.md': DocumentType.MD
        }
        return mapping.get(file_extension, DocumentType.TXT)

    async def _save_chunks(self, document_id: int, chunks: List[DocumentChunkData]):
        """Sauvegarder les chunks en base de données"""
        try:
            chunk_objects = []
            
            for chunk_data in chunks:
                chunk = DocumentChunk(
                    document_id=document_id,
                    content=chunk_data.content,
                    chunk_index=chunk_data.chunk_index,
                    chunk_size=len(chunk_data.content),
                    word_count=len(chunk_data.content.split()),
                    start_position=chunk_data.start_position,
                    end_position=chunk_data.end_position,
                    metadata=chunk_data.metadata
                )
                chunk_objects.append(chunk)
            
            self.db.add_all(chunk_objects)
            self.db.flush()  # Important: ne pas commit ici, laisser le contrôle à process_document
            
            logger.info(f"Sauvegarde de {len(chunk_objects)} chunks pour le document {document_id}")
            
        except Exception as e:
            logger.error(f"Erreur sauvegarde chunks: {str(e)}")
            raise

    # ✅ CORRECTION: Supprimer l'ancienne méthode _generate_embeddings pour éviter la confusion
    # Elle est remplacée par _generate_embeddings_sync

    def get_documents(
        self, 
        user_id: int, 
        limit: int = 20, 
        offset: int = 0,
        status_filter: Optional[EmbeddingStatus] = None
    ) -> List[Document]:
        """Récupérer les documents d'un utilisateur"""
        try:
            query = self.db.query(Document).filter(Document.uploaded_by == user_id)
            
            if status_filter:
                query = query.filter(Document.embeddings_status == status_filter)
            
            return query.order_by(desc(Document.created_at)).offset(offset).limit(limit).all()
            
        except Exception as e:
            logger.error(f"Erreur récupération documents pour utilisateur {user_id}: {str(e)}")
            return []

    def get_document(self, document_id: int, user_id: int) -> Optional[Document]:
        """Récupérer un document spécifique"""
        try:
            return self.db.query(Document).filter(
                and_(Document.id == document_id, Document.uploaded_by == user_id)
            ).first()
        except Exception as e:
            logger.error(f"Erreur récupération document {document_id}: {str(e)}")
            return None

    def delete_document(self, document_id: int, user_id: int) -> bool:
        """Supprimer un document"""
        try:
            document = self.get_document(document_id, user_id)
            if not document:
                logger.warning(f"Document {document_id} non trouvé pour utilisateur {user_id}")
                return False
            
            # Supprimer des embeddings si le service est disponible
            if self.embedding_service:
                try:
                    asyncio.create_task(
                        self.embedding_service.delete_document(str(document_id))
                    )
                except Exception as e:
                    logger.warning(f"Erreur suppression embeddings: {str(e)}")
            
            # Supprimer le document
            self.db.delete(document)
            self.db.commit()
            
            logger.info(f"Document {document_id} supprimé pour utilisateur {user_id}")
            return True
            
        except Exception as e:
            logger.error(f"Erreur suppression document {document_id}: {str(e)}")
            return False

    def get_document_stats(self, user_id: int) -> Dict[str, Any]:
        """Récupérer les statistiques de documents d'un utilisateur"""
        try:
            total_docs = self.db.query(Document).filter(Document.uploaded_by == user_id).count()
            
            processed_docs = self.db.query(Document).filter(
                and_(
                    Document.uploaded_by == user_id,
                    Document.embeddings_status == EmbeddingStatus.COMPLETED
                )
            ).count()
            
            pending_docs = self.db.query(Document).filter(
                and_(
                    Document.uploaded_by == user_id,
                    Document.embeddings_status == EmbeddingStatus.PENDING
                )
            ).count()
            
            processing_docs = self.db.query(Document).filter(
                and_(
                    Document.uploaded_by == user_id,
                    Document.embeddings_status == EmbeddingStatus.PROCESSING
                )
            ).count()
            
            return {
                "total_documents": total_docs,
                "processed_documents": processed_docs,
                "pending_documents": pending_docs,
                "processing_documents": processing_docs,
                "processing_rate": round(processed_docs / total_docs * 100, 2) if total_docs > 0 else 0
            }
            
        except Exception as e:
            logger.error(f"Erreur statistiques documents pour utilisateur {user_id}: {str(e)}")
            return {
                "total_documents": 0,
                "processed_documents": 0,
                "pending_documents": 0,
                "processing_documents": 0,
                "processing_rate": 0
            }